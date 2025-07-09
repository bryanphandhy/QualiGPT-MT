from flask import Flask, render_template, request, jsonify, send_file
import pandas as pd
# Keep optional OpenAI import for legacy compatibility; primary flow now uses llm_providers
# but importing it conditionally avoids breaking environments without the package.
try:
    from openai import OpenAI  # type: ignore
except ModuleNotFoundError:
    OpenAI = None  # type: ignore
import os
import io
import json
from werkzeug.utils import secure_filename
from docx import Document
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
import re
from datetime import datetime
import tempfile
from llm_providers import get_provider

# Download NLTK data if not already present
import nltk.data
# Set local NLTK data path
local_nltk_path = os.path.join(os.path.dirname(__file__), 'nltk_data')
if local_nltk_path not in nltk.data.path:
    nltk.data.path.append(local_nltk_path)

try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    # Download to local directory
    nltk.download('punkt_tab', download_dir=local_nltk_path)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Allowed file extensions
ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_participant_id(filename):
    """Extract participant ID from filename by removing file extension"""
    # Remove file extension and use the base filename as participant ID
    participant_id = filename.rsplit('.', 1)[0] if '.' in filename else filename
    # Clean up the participant ID (remove any special characters that might cause issues)
    participant_id = re.sub(r'[^\w\-_]', '', participant_id)
    return participant_id if participant_id else 'Unknown'

def add_participant_codes_to_content(content, participant_id):
    """Prefix every line of content with the participant code in square brackets so the LLM can attribute quotes accurately."""
    lines = content.split('\n')
    coded_lines = [f"[{participant_id}] {ln.strip()}" for ln in lines if ln.strip()]
    return '\n'.join(coded_lines)

# Prompt templates
PROMPTS = {
    "Interview": """You need to analyze an dataset of interviews.
Please identify the top {num_themes} key themes from the interview and organize the results in a structured table format.
The table should includes these items:
- 'Theme': Represents the main idea or topic identified from the interview.
- 'Description': Provides a brief explanation or summary of the theme.
- 'Quotes': Contains the complete, verbatim quotations from participants that fully capture their opinions and support the identified theme (do NOT truncate these quotes). IMPORTANT: Each quote MUST end with the participant code/ID in square brackets, for example: "This is what I think about the topic" [P001] (not parentheses or any other symbol).
- 'Participant Count': Indicates the number of participants who mentioned or alluded to the theme.
The table should be formatted as follows:
Each column should be separated by a '|' symbol, and there should be no extra '|' symbols within the data. Each row should end with '---'.
The whole table should start with '**********' and end with '**********'.
Columns: | 'Theme' | 'Description' | 'Quotes' | 'Participant Count' |.
Ensure each row of the table represents a distinct theme and its associated details. Aggregate the counts for each theme to show the total number of mentions across all participants.

IMPORTANT: Output ONLY the table with no additional text, commentary, or explanations. Do not include phrases like 'Here is the table', 'Below is the analysis', or 'Certainly!'. Start your response immediately with '**********' and end with '**********'. Do not use markdown formatting or code blocks.""",
    
    "Focus Group": """You need to analyze an dataset of a focus group.
Please identify the {num_themes} most common key themes from the interview and organize the results in a structured table format.
The table should includes these items:
- 'Theme': Represents the main idea or topic identified from the interview.
- 'Description': Provides a brief explanation or summary of the theme.
- 'Quotes': Contains the complete, verbatim quotations from participants that fully capture their opinions and support the identified theme (do NOT truncate these quotes). IMPORTANT: Each quote MUST end with the participant code/ID in square brackets, for example: "This is what I think about the topic" [P001] (not parentheses or any other symbol).
- 'Participant Count': Indicates the number of participants who mentioned or alluded to the theme. Please ensure this count reflects the actual number of participants who discussed each theme.
The table should be formatted strictly as follows:
The table should have 4 columns only.
Each column should be separated by a '|' symbol, and there should be no extra '|' symbols within the data. Each row should end with '---'.
Start the table with '**********'.
The header row should be: | 'Theme' | 'Description' | 'Quotes' | 'Participant Count' |
Followed by a row of '|---|---|---|---|'.
End the table with '**********'.
Each subsequent row should represent a theme and its details, with columns separated by '|'.
Ensure each row of the table represents a distinct theme and its associated details.

IMPORTANT: Output ONLY the table with no additional text, commentary, or explanations. Do not include phrases like 'Here is the table', 'Below is the analysis', or 'Certainly!'. Start your response immediately with '**********' and end with '**********'. Do not use markdown formatting or code blocks.""",
    
    "Social Media Posts": """You need to analyze an dataset of Social Media Posts.
Please identify the top {num_themes} key themes from the interview and organize the results in a structured table format.
The table should includes these items:
- 'Theme': Represents the main idea or topic identified from the interview.
- 'Description': Provides a brief explanation or summary of the theme.
- 'Quotes': Contains the complete, verbatim quotations from participants that fully capture their opinions and support the identified theme (do NOT truncate these quotes). IMPORTANT: Each quote MUST end with the participant code/ID in square brackets, for example: "This is what I think about the topic" [P001] (not parentheses or any other symbol).
- 'Participant Count': Indicates the number of participants who mentioned or alluded to the theme.
The table should be formatted as follows:
Each column should be separated by a '|' symbol, and there should be no extra '|' symbols within the data. Each row should end with '---'.
The whole table should start with '**********' and end with '**********'.
Columns: | 'Theme' | 'Description' | 'Quotes' | 'Participant Count' |.
Ensure each row of the table represents a distinct theme and its associated details.

IMPORTANT: Output ONLY the table with no additional text, commentary, or explanations. Do not include phrases like 'Here is the table', 'Below is the analysis', or 'Certainly!'. Start your response immediately with '**********' and end with '**********'. Do not use markdown formatting or code blocks."""
}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/test_api', methods=['POST'])
def test_api():
    try:
        data = request.json
        api_key = data.get('api_key')
        provider_name = data.get('provider', 'openai')
        model_name = data.get('model')  # Get the specific model
        
        if not api_key:
            return jsonify({'success': False, 'error': 'API key is required'})
        
        try:
            provider = get_provider(provider_name, api_key)
            
            # For Gemini, we can test with a specific model
            if provider_name == 'gemini' and model_name:
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(model_name)
                _ = model.generate_content("ping", generation_config={"max_output_tokens": 1})
            else:
                # For other providers, use default test_connection
                provider.test_connection()
                
        except Exception as e:
            return jsonify({'success': False, 'error': f'Connection failed: {e}'})
        
        return jsonify({'success': True, 'message': f'Successfully connected to {model_name or provider_name}'})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/upload_file', methods=['POST'])
def upload_file():
    try:
        if 'files' not in request.files:
            return jsonify({'success': False, 'error': 'No files uploaded'})
        
        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            return jsonify({'success': False, 'error': 'No files selected'})
        
        processed_files = []
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_ext = filename.rsplit('.', 1)[1].lower()
                
                if file_ext == 'csv':
                    data = pd.read_csv(file)
                elif file_ext == 'xlsx':
                    data = pd.read_excel(file)
                elif file_ext == 'docx':
                    def _extract_text_from_docx(file_obj):
                        """Return a list of non-empty text lines from a DOCX file, looking at paragraphs *and* table cells."""
                        doc = Document(file_obj)
                        lines: list[str] = []

                        for para in doc.paragraphs:
                            text = para.text.strip()
                            if text:
                                lines.append(text)

                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    cell_text = cell.text.strip()
                                    if cell_text:
                                        for piece in cell_text.split("\n"):
                                            piece = piece.strip()
                                            if piece:
                                                lines.append(piece)
                        return lines

                    full_text = _extract_text_from_docx(file)
                    if not full_text:
                        full_text = ["(No readable text detected in DOCX)"]
                    data = pd.DataFrame(full_text, columns=['Content'])
                
                headers = list(data.columns)
                data_content = '\n'.join(data.apply(lambda row: ' '.join(row.astype(str)), axis=1))
                
                # Extract participant ID from filename
                participant_id = extract_participant_id(filename)
                
                processed_files.append({
                    'filename': filename,
                    'participant_id': participant_id,
                    'headers': headers,
                    'data_content': data_content
                })

        if not processed_files:
            return jsonify({'success': False, 'error': 'Invalid file types or empty files'})

        return jsonify({
            'success': True,
            'files': processed_files
        })
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        data = request.json
        api_key = data.get('api_key')
        provider_name = data.get('provider', 'openai')
        model_name = data.get('model')
        
        # Handle single file vs multiple files
        files_data = data.get('files_data')
        analysis_mode = data.get('analysis_mode', 'combined')

        data_type = data.get('data_type')
        num_themes = data.get('num_themes', 10)
        custom_prompt = data.get('custom_prompt', '')
        enable_role_playing = data.get('enable_role_playing', False)
        pre_detect_themes = data.get('pre_detect_themes', False)
        temperature = data.get('temperature', 0.7)
        max_tokens = data.get('max_tokens', 4000)
        english_output = data.get('english_output', False)

        if not api_key or not files_data:
            return jsonify({'success': False, 'error': 'API key and data content are required'})

        provider = get_provider(provider_name, api_key)

        vietnamese_instruction = (
            " Nếu dữ liệu nguồn có vẻ được viết bằng tiếng Việt, hãy trình bày toàn bộ bảng (bao gồm tiêu đề cột, mô tả, trích dẫn) bằng tiếng Việt."
        )
        english_instruction = (
            " Present the entire table in English, with correct grammar and spelling, translating and grammar-correcting any participant quotes as needed."
        )
        language_instruction = english_instruction if english_output else vietnamese_instruction

        if enable_role_playing:
            system_message = (
                "You are an excellent qualitative data analyst and qualitative research expert. "
                "Follow the output format instructions exactly with no additional commentary." + language_instruction
            )
        else:
            system_message = (
                "You are a helpful assistant. Follow the output format instructions exactly with no additional commentary." + language_instruction
            )

        def _run_single_analysis(content, participant_id=None):
            # If num_themes is 'auto', prompt the LLM to choose the optimal number
            if num_themes == 'auto':
                prompt = (
                    "You need to analyze a dataset of interviews. "
                    "Identify the optimal number of key themes (no more than 20) that comprehensively cover all significant ideas and perspectives in the data. "
                    "Present a table of all major and minor themes, ensuring no important information is lost. "
                    "Do not limit the number of themes unless the data naturally supports fewer themes. "
                    "The table should include: | 'Theme' | 'Description' | 'Quotes' | 'Participant Count' |. "
                    "IMPORTANT: Output ONLY the table with no additional text, commentary, or explanations. Start your response immediately with '**********' and end with '**********'. Do not use markdown formatting or code blocks. "
                )
            else:
                if custom_prompt:
                    prompt = custom_prompt
                else:
                    prompt = PROMPTS.get(data_type, PROMPTS['Interview']).format(num_themes=num_themes)

            # Add participant code context if provided
            if participant_id:
                content = add_participant_codes_to_content(content, participant_id)

            segments = split_into_segments(content)
            all_responses = []
            for segment in segments:
                combined_message = segment + "\n\n" + prompt
                response_text = provider.chat(
                    system_message,
                    combined_message,
                    model=model_name or "auto",
                    temperature=temperature,
                    max_tokens=max_tokens,
                )
                all_responses.append(response_text)

            if len(segments) > 1:
                merged_responses = "\n".join(all_responses)
                return analyze_merged_responses(
                    merged_responses, num_themes, system_message, provider, model_name, temperature, max_tokens
                )
            else:
                # Fallback: If auto mode and output is empty or malformed, retry with num_themes=10
                if num_themes == 'auto':
                    parsed = parse_response_to_csv(all_responses[0])
                    if not parsed or len(parsed) < 2:
                        fallback_prompt = PROMPTS.get(data_type, PROMPTS['Interview']).format(num_themes=10)
                        fallback_message = segment + "\n\n" + fallback_prompt
                        fallback_response = provider.chat(
                            system_message,
                            fallback_message,
                            model=model_name or "auto",
                            temperature=temperature,
                            max_tokens=max_tokens,
                        )
                        return fallback_response
                return all_responses[0]

        if analysis_mode == 'combined':
            # For combined analysis, include participant IDs in the content
            combined_content_parts = []
            for f in files_data:
                participant_content = add_participant_codes_to_content(f['data_content'], f['participant_id'])
                combined_content_parts.append(participant_content)
            combined_content = "\n\n".join(combined_content_parts)
            
            final_response = _run_single_analysis(combined_content)
            # Check for empty or malformed output
            parsed = parse_response_to_csv(final_response)
            if not parsed or len(parsed) < 2:
                return jsonify({'success': False, 'error': 'AI did not return a valid table. Try reducing the number of files, or use a fixed number of themes.'})
            # If auto, count number of themes in the table
            num_themes_auto = None
            if num_themes == 'auto':
                num_themes_auto = len(parsed) - 1
            return jsonify({
                'success': True,
                'response': final_response,
                'report_type': 'combined',
                'segments_processed': len(split_into_segments(combined_content)),
                'num_themes_auto': num_themes_auto
            })
        else: # separate reports
            separate_results = []
            for file_data in files_data:
                analysis_result = _run_single_analysis(file_data['data_content'], file_data['participant_id'])
                parsed = parse_response_to_csv(analysis_result)
                if not parsed or len(parsed) < 2:
                    return jsonify({'success': False, 'error': f"AI did not return a valid table for file {file_data['filename']}. Try using a fixed number of themes or fewer files."})
                num_themes_auto = None
                if num_themes == 'auto':
                    num_themes_auto = len(parsed) - 1
                separate_results.append({
                    'filename': file_data['filename'],
                    'participant_id': file_data['participant_id'],
                    'analysis': analysis_result,
                    'num_themes_auto': num_themes_auto
                })
            
            return jsonify({
                'success': True,
                'response': separate_results,
                'report_type': 'separate'
            })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

def split_into_segments(text, max_tokens=120000):
    """Split text into segments that fit within GPT-4o's token limits
    
    GPT-4o has 128k context window, so we use 120k for data and reserve 8k for prompts/responses.
    This is a ~30x increase from the previous 3800 token limit for GPT-3.5-turbo.
    Most datasets will now process in a single call.
    """
    try:
        sentences = sent_tokenize(text)
    except:
        # Fallback to simple splitting if NLTK fails
        sentences = text.split('.')
        sentences = [s + '.' for s in sentences if s.strip()]
    
    segments = []
    segment = ""
    segment_tokens = 0
    
    for sentence in sentences:
        num_tokens = len(word_tokenize(sentence)) if 'word_tokenize' in globals() else len(sentence.split())
        
        if segment_tokens + num_tokens > max_tokens:
            segments.append(segment.strip())
            segment = sentence
            segment_tokens = num_tokens
        else:
            segment += " " + sentence
            segment_tokens += num_tokens
    
    if segment:
        segments.append(segment.strip())
    
    return segments

def analyze_merged_responses(merged_responses, num_themes, system_message, provider, model_name, temperature, max_tokens):
    """Analyze merged responses to create a final summary"""
    prompt = f"""This is the result of a thematic analysis of several parts of the dataset. Now, summarize the same themes to generate a new table.
Please identify the {num_themes} most common key themes from the interview and organize the results in a structured table format.
The table should include the following columns:
'Theme': Represents the main idea or topic identified from the interview.
'Description': Provides a brief explanation or summary of the theme.
'Quotes': Contains the complete, verbatim quotations from participants that fully capture their opinions and support the identified theme (do NOT truncate these quotes). IMPORTANT: Each quote MUST end with the participant code/ID in square brackets, for example: "This is what I think about the topic" [P001] (not parentheses or any other symbol).
'Participant Count': Indicates the number of participants who mentioned or alluded to the theme.
The table should be formatted strictly as follows:
- Start the table with '**********'.
- The header row should be: | 'Theme' | 'Description' | 'Quotes' | 'Participant Count' |
- Followed by a row of '|---|---|---|---|'.
- Each subsequent row should represent a theme and its details, with columns separated by '|'.
- Each row should end with '---'.
- End the table with '**********'.
Ensure each row of the table represents a distinct theme and its associated details.

IMPORTANT: Output ONLY the table with no additional text, commentary, or explanations. Do not include phrases like 'Here is the table', 'Below is the analysis', or 'Certainly!'. Start your response immediately with '**********' and end with '**********'. Do not use markdown formatting or code blocks.

Analyze the following merged responses: {merged_responses}"""
    
    response_text = provider.chat(
        system_message,
        prompt,
        model=model_name or "auto",  # Use selected model or default
        temperature=temperature,
        max_tokens=max_tokens,
    )
    
    return response_text

@app.route('/export_csv', methods=['POST'])
def export_csv():
    try:
        data = request.json
        response_content = data.get('response', '')
        
        # Parse the response to extract table data
        parsed_data = parse_response_to_csv(response_content)
        
        if not parsed_data:
            return jsonify({'success': False, 'error': 'Failed to parse the response'})
        
        # Create DataFrame
        df = pd.DataFrame(parsed_data[1:], columns=parsed_data[0])
        
        # Create CSV in memory
        output = io.StringIO()
        df.to_csv(output, index=False)
        output.seek(0)
        
        # Create response
        return send_file(
            io.BytesIO(output.getvalue().encode()),
            mimetype='text/csv',
            as_attachment=True,
            download_name=f'qualigpt_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
        )
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

def parse_response_to_csv(response):
    """Parse the GPT response to extract table data"""
    lines = response.strip().split("\n")
    
    # Find table delimiters
    delimiter_indices = [i for i, line in enumerate(lines) if line.strip() == "**********"]
    
    if len(delimiter_indices) < 2:
        return []
    
    start_index, end_index = delimiter_indices[0], delimiter_indices[-1]
    table_content = lines[start_index+1:end_index]
    
    # Parse table rows
    parsed_data = []
    for line in table_content:
        if line.strip() and '|' in line and not line.strip().startswith('|---'):
            # Split by | and clean up
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty cells at start and end
            cells = [cell for cell in cells if cell]
            if cells:
                parsed_data.append(cells)
    
    return parsed_data

if __name__ == '__main__':
    app.run(debug=True, port=5000)
